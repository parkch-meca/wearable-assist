"""국문 학술지 논문용 docx 생성 도구 모음 (python-docx).

pandoc 미설치 + 캡션 위치·네이티브 표·한글 폰트·줄간격·페이지 번호를
결정적으로 제어해야 하므로 python-docx 직접 생성 방식을 사용한다.

제공 기능
  set_korean_font  : 라틴/한글(eastAsia) 폰트를 함께 지정 (한글 폰트 미지정 시 Word가 임의 대체)
  add_toc          : 실제 Word TOC 필드 + 캐시된 렌더 결과 (Word는 갱신 가능, PDF 변환은 캐시를 렌더)
  add_page_number  : 바닥글 PAGE 필드
  add_table        : 네이티브 표(이미지 아님) + 표 위 캡션
  add_figure       : 이미지 임베드 + 그림 아래 캡션
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = 'Malgun Gothic'      # 윈도우 기본 한글 폰트
LATIN_FONT = 'Times New Roman'
BODY_PT = 10.5
LINE = 1.7                        # 170 %


def set_korean_font(run, size=None, name=None, latin=None, bold=None, italic=None,
                    color=None):
    """한글은 w:eastAsia 속성으로 별도 지정해야 적용된다."""
    name = name or BODY_FONT
    latin = latin or name
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    rFonts.set(qn('w:eastAsia'), name)
    return run


def para(doc, text='', size=BODY_PT, bold=False, align=None, line=LINE,
         space_before=0, space_after=4, name=None, latin=None, italic=False,
         color=None, indent_first=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent_first is not None:
        pf.first_line_indent = Cm(indent_first)
    if text:
        set_korean_font(p.add_run(text), size=size, bold=bold, name=name,
                        latin=latin, italic=italic, color=color)
    return p


def rich_para(doc, parts, size=BODY_PT, align=None, line=LINE, space_before=0,
              space_after=4, indent_first=None):
    """parts = [(text, {'bold':True, 'size':9, ...}), ...]"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent_first is not None:
        pf.first_line_indent = Cm(indent_first)
    for txt, kw in parts:
        kw = dict(kw or {})
        set_korean_font(p.add_run(txt), size=kw.pop('size', size), **kw)
    return p


def heading(doc, text, level=1, size=None, space_before=14, space_after=6):
    """개요 수준을 부여해 TOC 필드가 잡을 수 있게 한다."""
    size = size or {1: 13, 2: 11.5, 3: 10.5}[level]
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    pf = p.paragraph_format
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    pf.line_spacing = 1.3
    set_korean_font(p.add_run(text), size=size, bold=True, color=(0, 0, 0))
    return p


def _fld(p, instr, cached=None):
    r = p.add_run()
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin')
    r._element.append(fc)
    r2 = p.add_run()
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = instr
    r2._element.append(it)
    r3 = p.add_run()
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    r3._element.append(sep)
    if cached is not None:
        set_korean_font(p.add_run(cached), size=BODY_PT)
    r5 = p.add_run()
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    r5._element.append(end)


def add_toc(doc, entries, page_nums=None):
    """TOC 필드 + 캐시 결과.
    Word는 F9로 갱신 가능하고, LibreOffice PDF 변환은 캐시된 항목을 그대로 렌더한다.
    entries = [(제목, 수준)] — 캐시 표시용.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin')
    fc.set(qn('w:dirty'), 'true')
    r._element.append(fc)
    r2 = p.add_run()
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = r'TOC \o "1-2" \h \z \u'
    r2._element.append(it)
    r3 = p.add_run()
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    r3._element.append(sep)
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    for idx, (title, lvl) in enumerate(entries):
        q = doc.add_paragraph()
        qf = q.paragraph_format
        qf.line_spacing = 1.3
        qf.space_before = Pt(0); qf.space_after = Pt(1)
        qf.left_indent = Cm(0.0 if lvl == 1 else 0.7)
        qf.tab_stops.add_tab_stop(Cm(16.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        txt = title
        if page_nums:
            txt = f'{title}\t{page_nums[idx]}'
        set_korean_font(q.add_run(txt), size=10 if lvl == 1 else 9.5,
                        bold=(lvl == 1))
    last = doc.add_paragraph()
    last.paragraph_format.space_after = Pt(0)
    r5 = last.add_run()
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    r5._element.append(end)


def add_page_number(doc):
    sec = doc.sections[0]
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _fld(p, 'PAGE', cached='1')


def add_caption(doc, text, above=True, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.3
    pf.space_before = Pt(8 if above else 4)
    pf.space_after = Pt(4 if above else 10)
    pf.keep_with_next = above          # 표/그림과 캡션이 분리되지 않게
    p.alignment = align
    # "Table 1." / "Figure 1." 부분만 굵게
    if '. ' in text:
        head, rest = text.split('. ', 1)
        set_korean_font(p.add_run(head + '. '), size=size, bold=True)
        set_korean_font(p.add_run(rest), size=size)
    else:
        set_korean_font(p.add_run(text), size=size, bold=True)
    return p


def add_table(doc, caption, header, rows, widths=None, size=9.0,
              align_right_cols=(), keep_together=True):
    """네이티브 Word 표 + 표 위 캡션."""
    add_caption(doc, caption, above=True)
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0]
    for i, h in enumerate(header):
        c = hdr.cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(1)
        set_korean_font(p.add_run(str(h)), size=size, bold=True)
        _shade(c, 'E8ECEF')
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(1)
            p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols
                           else (WD_ALIGN_PARAGRAPH.CENTER if i == 0 and False
                                 else WD_ALIGN_PARAGRAPH.LEFT))
            txt = str(v)
            bold = txt.startswith('**') and txt.endswith('**')
            if bold:
                txt = txt.strip('*')
            set_korean_font(p.add_run(txt), size=size, bold=bold)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    # 머리행 반복 (페이지 넘김 시)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    hdr_el = OxmlElement('w:tblHeader'); hdr_el.set(qn('w:val'), 'true')
    trPr.append(hdr_el)
    # 행 중간 분할 금지
    for r in t.rows:
        trPr = r._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit'); trPr.append(cs)
    # 표 전체를 한 페이지에 유지 (마지막 행 제외 전 행에 keep_with_next)
    if keep_together:
        for r in list(t.rows)[:-1]:
            for c in r.cells:
                for pp in c.paragraphs:
                    pp.paragraph_format.keep_with_next = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def add_figure(doc, path, caption, width_cm=16.0):
    """이미지를 문서에 임베드(경로 링크 아님) + 그림 아래 캡션."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(path, width=Cm(width_cm))
    add_caption(doc, caption, above=False, align=WD_ALIGN_PARAGRAPH.LEFT)


def new_doc(margins_cm=2.2):
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21.0)      # A4 (국문 학술지 표준)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(margins_cm); s.bottom_margin = Cm(margins_cm)
        s.left_margin = Cm(margins_cm); s.right_margin = Cm(margins_cm)
    st = doc.styles['Normal']
    st.font.name = LATIN_FONT
    st.font.size = Pt(BODY_PT)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    return doc
