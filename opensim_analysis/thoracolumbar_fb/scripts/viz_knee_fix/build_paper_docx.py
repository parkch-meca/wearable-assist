"""논문 완성본(docx + pdf) 생성 — 그림 임베드 + 네이티브 표.

pandoc 미설치이고 캡션 위치·네이티브 표·한글 폰트·줄간격·페이지 번호를 결정적으로
제어해야 하므로 python-docx 직접 생성 방식을 사용한다 (docx_kit.py).
수치는 paper_numbers 모듈(단일 소스)을 경유해서만 읽는다. 직접 JSON 로드 금지.
"""
import json, os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx_kit import (new_doc, para, rich_para, heading, add_toc, add_page_number,
                      add_table, add_figure, add_caption, set_korean_font, BODY_PT)
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
from docx.shared import Pt, Cm
import paper_numbers as pn      # 표기 규칙 단일 소스 — 직접 반올림 금지
import docx_kit
docx_kit.LINE = 1.6           # 160 % — 그림 배치 여지 확보 (규격 160~180 % 내)

FIG = '/data/wearable-assist/opensim_analysis/thoracolumbar_fb/docs/images/paper_five_motion'
OUT = '/data/wearable-assist/opensim_analysis/thoracolumbar_fb/docs/five_motion_paper.docx'
U, G = pn.U, pn.G          # 경로 분기 방지 — paper_numbers 가 유일한 로더
MA = {k: pn.metric(k, 'a') for k in pn.ORDER}   # A/B/C 는 이미 쓰이는 이름
MB = {k: pn.metric(k, 'b') for k in pn.ORDER}
MC = {k: pn.metric(k, 'c') for k in pn.ORDER}
NM = {'squat': '맨몸 스쿼트', 'stoop': '맨몸 스툽', 'box': '박스 들기',
      'gait': '맨몸 보행', 'carry': '박스 운반'}
LD = {'squat': '0 kg', 'stoop': '0 kg', 'box': '20 kg', 'gait': '0 kg', 'carry': '20 kg'}
ORD = ['squat', 'stoop', 'box', 'gait', 'carry']   # 전 표 공통 순서


def f(x, d=2):
    return f'{x:.{d}f}'.replace('-', '\u2212')      # 음수 기호를 본문과 통일


# 목차 쪽번호: 2-pass. 1차 실행 시 비어 있고, 생성된 PDF에서 추출해 재실행하면 채워진다.
_PN = os.path.join(os.path.dirname(OUT), '.toc_pages.json')
PAGE_NUMS = json.load(open(_PN)) if os.path.exists(_PN) else None


doc = new_doc()

# ============================================================ 표제
para(doc, 'SMA 직물 근육 액추에이터 기반 허리 보조 슈트의', size=15, bold=True,
     align=A.CENTER, line=1.35, space_after=2)
para(doc, '작업 동작별 척추기립근 부하 저감 효과', size=15, bold=True,
     align=A.CENTER, line=1.35, space_after=2)
para(doc, '— 전신 근골격계 시뮬레이션 기반 5동작 평가 —', size=12.5, bold=True,
     align=A.CENTER, line=1.35, space_after=12)
para(doc, '박철훈', size=11, align=A.CENTER, space_after=2)
para(doc, '한국기계연구원 (KIMM)', size=10, align=A.CENTER, space_after=2)
para(doc, 'parkch@kimm.re.kr', size=9.5, align=A.CENTER, space_after=16)

# ============================================================ 국문 초록
heading(doc, '국문 초록', 1, size=12, space_before=4, space_after=6)
para(doc, '작업 관련 요통을 줄이기 위한 능동 웨어러블 슈트가 다수 개발되고 있으나, 작업 종류에 따라 '
          '효과가 어떻게 달라지는지에 대한 정량 근거는 부족하다. 본 연구는 형상기억합금(SMA) 직물 근육 '
          '액추에이터 기반 허리 보조 슈트를 대상으로, 부하 조건이 서로 다른 5개 작업 동작(맨몸 스쿼트, '
          '맨몸 스툽, 20 kg 박스 들기, 맨몸 보행, 20 kg 박스 운반)에서 척추기립근(ES) 활성도를 정량 '
          '평가하였다. ThoracolumbarFB v2.0 전신 모델(620 근육, ES 76개)과 OpenSim Static Optimization을 '
          '사용하였고, 슈트는 흉추1–골반 간 순수 토크 커플 24 N·m으로 모델링하였다. 5개 동작은 완전히 '
          '동일한 모델 파일과 해석 설정을 공유하며, 이는 파일 해시로 검증하였다.',
     size=9.5, line=1.55, align=A.JUSTIFY, indent_first=0.5, space_after=5)
para(doc, '주 지표로 슈트 토크가 최대치의 90 % 이상인 구간의 ES peak 평균을 사전 정의하였다. 그 결과 '
          f'맨몸 스쿼트 {MB["squat"]["rel_s"]} %, 맨몸 스툽 {MB["stoop"]["rel_s"]} %, '
          f'박스 들기 {MB["box"]["rel_s"]} %, 박스 운반 {MB["carry"]["rel_s"]} %의 활성도 감소가 '
          '관찰되었다. 부하가 클수록 상대 감소율이 작아지는 경향이 나타났으나, 이 경향은 지표 정의에 따라 '
          '성립하지 않았다. 맨몸 보행에서는 감소가 아니라 재분배가 관찰되었다. 전체 활성도 총량은 '
          '−17.4 %p 감소하는 반면 최대 활성 근육은 +21.4 % 증가하며, 장늑근(Iliocostalis)이 −90.9 % '
          '비활성화되는 대신 최장근 요추부(Longissimus pars lumborum)가 +28.5 % 증가하였다. 방법론적으로, '
          '저부하 동작에서 reserve 액추에이터 설정이 근육 활성도를 3.2배 과소평가하고 슈트 효과의 존재 '
          '여부까지 왜곡할 수 있음을 제시하였다. 통일 조건의 스툽 미착용 활성도(68.9 %)는 선행 실측 '
          'EMG 연구의 69.8 %MVC와 0.9 %p 차이로 정합하였다.',
     size=9.5, line=1.55, align=A.JUSTIFY, indent_first=0.5, space_after=6)
rich_para(doc, [('핵심어: ', {'bold': True}),
                ('웨어러블 로봇, 형상기억합금, 척추기립근, 근골격계 시뮬레이션, '
                 'OpenSim, 정적 최적화, 요통 예방', {})],
          size=9.5, line=1.5, space_after=14)

# ============================================================ 영문 초록
heading(doc, 'Abstract', 1, size=12, space_before=4, space_after=6)
para(doc, 'Active wearable suits for occupational low-back pain are being developed widely, but '
          'quantitative evidence on how their effect varies across work tasks remains limited. This '
          'study quantified erector spinae (ES) activation under a shape-memory-alloy (SMA) fabric '
          'muscle lumbar-assist suit across five work tasks spanning a wide load range: bodyweight '
          'squat, bodyweight stoop, 20 kg box lifting, level walking, and 20 kg box carrying. A '
          'full-body ThoracolumbarFB v2.0 model (620 muscles, 76 ES fascicles) was analysed with '
          'OpenSim Static Optimization; the suit was modelled as a pure torque couple of 24 N·m '
          'between thoracic1 and the pelvis. All five tasks share an identical model file and solver '
          'configuration, verified by file hash.',
     size=9.5, line=1.55, align=A.JUSTIFY, indent_first=0.5, space_after=5,
     name='Times New Roman', latin='Times New Roman')
para(doc, 'The primary outcome, pre-specified as the mean ES peak activation over the window in which '
          f'suit torque exceeds 90 % of its maximum, decreased by {abs(MB["squat"]["rel"]):.1f} % (squat), '
          f'{abs(MB["stoop"]["rel"]):.1f} % (stoop), {abs(MB["box"]["rel"]):.1f} % (lifting) and '
          f'{abs(MB["carry"]["rel"]):.1f} % (carrying). A monotonic decrease of relative benefit with '
          'increasing load was observed for the primary outcome and for the whole-cycle peak, but '
          'not for the mean-activation definition. During level walking the suit redistributed rather than reduced load: '
          'total activation fell by 17.4 percentage points while the most-active muscle rose by 21.4 %, '
          'with iliocostalis nearly silenced (−90.9 %) and longissimus pars lumborum increased '
          '(+28.5 %). Methodologically, standard reserve-actuator settings underestimated ES activation '
          'by a factor of 3.2 in low-load tasks and produced a spurious assistive effect. Under the '
          'unified settings, unassisted stoop activation (68.9 %) agreed with a published surface-EMG '
          'value of 69.8 % MVC to within 0.9 percentage points.',
     size=9.5, line=1.55, align=A.JUSTIFY, indent_first=0.5, space_after=6,
     name='Times New Roman', latin='Times New Roman')
rich_para(doc, [('Keywords: ', {'bold': True, 'latin': 'Times New Roman'}),
                ('wearable robot, shape memory alloy, erector spinae, musculoskeletal simulation, '
                 'OpenSim, static optimization, low back pain prevention',
                 {'latin': 'Times New Roman'})],
          size=9.5, line=1.45, space_after=10)

# ============================================================ 목차
doc.add_page_break()
heading(doc, '목    차', 1, size=12, space_before=0, space_after=6)
TOC = [('1. 서론', 1), ('1.1 기존 연구와의 관계', 2), ('1.2 기여', 2),
       ('2. 재료 및 방법', 1), ('2.1 근골격계 모델 — 5동작 완전 통일', 2),
       ('2.2 모델 정합성 수정 및 회귀 검증', 2), ('2.3 슈트 모델링', 2),
       ('2.4 동작 데이터', 2), ('2.5 외력 및 지면반력', 2),
       ('2.6 Static Optimization 및 reserve 액추에이터', 2), ('2.7 평가 지표', 2),
       ('2.8 동작 검증 절차', 2), ('2.9 해석 조건 전수 감사 및 폐기된 서술', 2),
       ('3. 결과', 1), ('3.1 동작별 ES 활성도 및 슈트 효과', 2),
       ('3.2 부하–효과 관계', 2), ('3.3 절대 활성도 수준', 2),
       ('3.4 맨몸 보행 — 감소가 아니라 재분배', 2),
       ('3.5 주 지표 선택이 결론에 미치는 영향', 2),
       ('3.6 Reserve 설정 민감도', 2),
       ('4. 선행 연구 대조', 1), ('5. 고찰', 1), ('6. 한계', 1),
       ('7. 향후 과제', 1), ('8. 결론', 1), ('참고문헌', 1)]
add_toc(doc, TOC, page_nums=PAGE_NUMS)
add_page_number(doc)
doc.add_page_break()

# ============================================================ 1. 서론
heading(doc, '1. 서론', 1)
para(doc, '작업 관련 요통은 산업재해의 주요 원인이며, 물건 들기·운반은 요추 부하를 유발하는 대표적 원인 '
          '작업이다. 고령 근로자와 간병 노동의 증가로 근골격계 부담은 확대되고 있고, 이에 대응하는 능동 '
          '보조 웨어러블이 다수 개발되고 있으나 "어떤 작업에서 얼마나 도움이 되는가"에 대한 정량 근거는 '
          '상대적으로 부족하다.', indent_first=0.5, align=A.JUSTIFY)
para(doc, '표면 근전도(EMG)는 착용 효과 평가의 표준 수단이지만 표층 근육만 측정 가능하고, 전극 위치와 '
          '개인차에 민감하며, 동일 동작에서 착용/미착용을 동시에 비교할 수 없고, 척추 내부 부하를 직접 '
          '추정하기 어렵다. 근골격계 시뮬레이션은 이 한계를 보완하여 심부 근육을 포함한 개별 근육 '
          '활성도를 동일 운동학 조건에서 슈트 ON/OFF로 직접 비교할 수 있게 한다.',
     indent_first=0.5, align=A.JUSTIFY)
para(doc, '본 연구는 SMA 직물 근육 액추에이터 기반 허리 보조 슈트를 대상으로, 부하 조건이 서로 다른 '
          '5개 작업 동작에서 완전히 동일한 모델·해석 조건으로 척추기립근(erector spinae, 이하 ES) '
          '활성도를 비교하여 슈트 효과의 작업 의존성을 정량화하는 것을 목적으로 한다.',
     indent_first=0.5, align=A.JUSTIFY)

heading(doc, '1.1 기존 연구와의 관계', 2)
para(doc, '본 연구진은 앞서 동일 슈트를 대상으로 stoop 단일 동작에 대해 MocoInverse 기반 도즈–반응 '
          '분석을 별도로 수행하였다. 해당 연구는 해석기(MocoInverse), 근육 set(척추 관련 114개), '
          '주 지표(ES mean), reserve 설정이 본 연구와 모두 다르므로 두 연구의 stoop 수치는 직접 '
          '등치될 수 없다. 본 연구는 stoop을 5동작 비교의 한 조건으로만 사용한다.',
     indent_first=0.5, align=A.JUSTIFY)

heading(doc, '1.2 기여', 2)
for t in ['부하 스펙트럼을 포괄하는 5개 작업 동작에서 완전히 동일한 모델 파일과 해석 설정으로 ES 활성도를 비교하였다.',
          '평가 지표 정의가 결론을 좌우함을 정량적으로 제시하고 지표 선택 기준을 제안하였다(3.5절).',
          '저부하 동작에서 reserve 액추에이터 설정이 근육 활성도를 3.2배 과소평가하고 슈트 효과의 부호까지 왜곡할 수 있음을 제시하였다(3.6절).',
          '정상 보행에서 슈트가 부하를 줄이는 것이 아니라 근육군 사이에 재분배함을 발견하고, 상시 구동 운용의 트레이드오프 가설을 제기하였다(3.4절, 5.3절).']:
    para(doc, '· ' + t, align=A.JUSTIFY, space_after=3)

# ============================================================ 2. 방법
heading(doc, '2. 재료 및 방법', 1)
heading(doc, '2.1 근골격계 모델 — 5동작 완전 통일', 2)
para(doc, 'ThoracolumbarFB v2.0 Full Body 모델[3]을 OpenSim 4.6에서 사용하였다. 전신 620개 근육, '
          '78개 body, 29개 joint로 구성되며 흉요추 T1–L5 전 분절이 굴곡·신전 자유도를 갖는다. '
          '5개 동작 전체가 단일 모델 파일을 공유하며, 이를 파일 해시(SHA-1 앞 12자리)로 검증하였다(Table 1).',
     indent_first=0.5, align=A.JUSTIFY)
add_table(doc, 'Table 1. 5동작 해석 조건 동일성 검증.',
          ['항목', '5동작 공통값', '검증'],
          [['모델 파일 전체 해시', 'ca12f321326e', '완전 동일'],
           ['기저 모델 해시 (reserve 제외)', 'e5bb8ab98934', '완전 동일'],
           ['좌표 수 / 근육 수 / reserve 수', '169 / 620 / 169', '동일'],
           ['흉쇄관절 2-DOF (M1 견갑)', '적용', '동일'],
           ['CoordinateCouplerConstraint', '0개 (제거)', '동일'],
           ['척추 reserve optimal_force', '5 N·m', '동일'],
           ['골반 reserve', '병진 500 N / 회전 1000 N·m', '동일'],
           ['기타 reserve', '병진 100 N / 회전 1000 N·m', '동일'],
           ['activation exponent / use_muscle_physiology', '2 / true', '동일']],
          widths=[6.4, 6.2, 3.4])
para(doc, '정량 대상 ES는 장늑근(IL, 24개), 최장근 요추부(LTpL, 10개), 흉추부(LTpT, 42개)의 76개이며, '
          '5동작에서 근육 목록이 원소 단위로 동일함을 확인하였다. 피험자 조건은 성인 남성(신장 약 '
          '175 cm, 체중 약 75 kg) 단일 체형이다. 모델과 슈트 모델링 개요를 Figure 1에 보인다.',
     indent_first=0.5, align=A.JUSTIFY)
add_figure(doc, f'{FIG}/fig1_model_and_suit.png',
           'Figure 1. 근골격계 모델과 슈트 모델링. (a) 전신 모델과 슈트 인장 경로, '
           '(b) 흉추1–골반 간 순수 토크 커플, (c) 액추에이터 사양에서 해석 토크 유도.', width_cm=16.4)

heading(doc, '2.2 모델 정합성 수정 및 회귀 검증', 2)
para(doc, '해석 과정에서 모델 자체의 정의 오류로 사람이 통상 수행하는 자세가 구현되지 않는 문제가 '
          '발견되었다. 자세를 임의로 왜곡해 맞추는 대신 모델 정의를 수정하고, 수정이 ES 정량에 미치는 '
          '영향을 회귀 검증으로 확인하였다(Table 2).', indent_first=0.5, align=A.JUSTIFY)
add_table(doc, 'Table 2. 모델 정합성 수정 항목 및 ES 영향.',
          ['수정 항목', '성격', 'ES 정량 영향'],
          [['전완 정의 오류 패치 폐기', '시각화 전용', '영향 없음'],
           ['손목 3자유도 잠금 해제', '시각화 전용', '영향 없음'],
           ['어깨대 좌우 대칭 처리', '시각화 전용', '영향 없음'],
           ['흉쇄관절 2-DOF 추가 (M1 견갑)', '모델 구조', 'ΔES 0.029 %p'],
           ['좌측 팔 관절축 7개 거울 대칭 수정', '모델 구조', 'ΔES ≤ 1.1 %p'],
           ['CoordinateCouplerConstraint 4개 제거', '모델 구조', 'ΔES ≤ 1.16 %p (선행 회귀 검증 유계)']],
          widths=[7.0, 4.0, 5.0])

heading(doc, '2.3 슈트 모델링', 2)
para(doc, '슈트는 흉추 1번(thoracic1)과 골반(pelvis) 사이에 작용하는 순수 토크 커플로 모델링하였다. '
          '흉추 분절에 신전 방향 +24 N·m, 골반에 −24 N·m을 부여하며, 모멘트 암의 해부학적 배치 가정에 '
          '의존하지 않는 보수적 표현이다. 5동작 모두 부착 body·성분·크기가 동일함을 외력 데이터에서 '
          '확인하였다. 이 값은 SMA 직물 근육의 편측 최대 수축력 100 N(양측 200 N)과 모멘트 암 '
          '0.10–0.13 m에서 유도한 20–26 N·m 범위의 대표값이다. 슈트 조건 외 모든 입력은 OFF/ON에서 '
          '동일하며, 근육 파라미터 조작은 수행하지 않았다.', indent_first=0.5, align=A.JUSTIFY)

heading(doc, '2.4 동작 데이터', 2)
add_table(doc, 'Table 3. 5개 작업 동작의 해석 조건.',
          ['동작', '외부 하중', '자세 특징', '운동학 출처', '지면반력', '프레임', '샘플링'],
          [['맨몸 스쿼트', '0 kg', '요추 중립, 무릎 굴곡 하강', '합성', '합성 GRF', '151', '30 Hz'],
           ['맨몸 스툽', '0 kg', '고관절 힌지, 요추 굴곡', '합성', '합성 GRF', '135', '26.7 Hz (비균일)'],
           ['박스 들기', '20 kg', '스툽 자세, 테이블 30 cm', '합성', '없음 (골반 reserve)', '226', '30 Hz'],
           ['맨몸 보행', '0 kg', '정상 보행 1주기', '실측 리타겟', '실측 GRF', '73', '60 Hz'],
           ['박스 운반', '20 kg', '전방 하중, 체간 후경 5°', '실측 리타겟', '실측 GRF + 하중 분배', '73', '60 Hz']],
          widths=[2.5, 1.7, 3.7, 2.1, 3.0, 1.3, 2.1], size=8.4)
para(doc, '합성 동작은 문헌 기반 관절 각도 목표를 설정하고 역기구학으로 생성하였다. 보행·운반은 '
          'gait2354 실측 역기구학 결과를 모델 좌표계로 리타겟하였다. 저역통과 필터는 합성 동작에 '
          '적용하지 않았고 실측 리타겟에는 6 Hz를 적용하였다. 합성 궤적은 매끄러워 필터가 불필요하고 '
          '실측 리타겟은 미분 잡음 억제가 필요하기 때문이며, 이는 의도된 차이이다. 5개 동작의 대표 '
          '자세는 Figure 2와 같다. 동작별 해석 조건은 Table 3에 정리하였다.',
     indent_first=0.5, align=A.JUSTIFY)
add_figure(doc, f'{FIG}/fig3_five_motion_postures.png',
           'Figure 2. 5개 작업 동작의 대표 자세. 각 패널은 슈트 미착용(좌)과 착용(우)을 나란히 보인다.',
           width_cm=16.4)

heading(doc, '2.5 외력 및 지면반력', 2)
para(doc, '박스 20 kg(196.2 N)은 양손에 균등 분배한 하방 외력(98.1 N/hand)으로 부여하였다. 보행·운반은 '
          '실측 지면반력을 사용하고 원 피험자와 모델의 체중 차이를 스케일로 보정하였으며, 압력중심을 '
          '모델 발 접지 위치에 정렬하였다. 운반은 박스 하중이 지면에도 전달되므로 지면반력 수직 성분에 '
          '발 접지 비율로 분배하여 중복 계산을 방지하였다. 박스 들기는 지면반력을 명시적으로 부여하지 '
          '않았고 골반 자유도 reserve가 체중 지지력을 흡수한다. 이는 OFF/ON에서 동일하므로 슈트 효과 '
          '차이에는 영향이 없으나, 박스 들기의 절대 활성도를 다른 동작과 직접 비교할 때는 고려해야 한다.',
     indent_first=0.5, align=A.JUSTIFY)

heading(doc, '2.6 Static Optimization 및 reserve 액추에이터', 2)
para(doc, '각 동작에 대해 슈트 토크만 0 / 24 N·m로 바꾸어 SO를 2회 실행하였다(총 10회). 척추 reserve의 '
          'optimal_force를 5 N·m로 낮춘 tight 설정을 5동작 전체에 적용하였다. optimal_force를 낮추면 '
          '같은 힘을 내기 위해 더 큰 control이 필요하고, SO의 비용함수(activation 제곱)가 이를 비싸게 '
          '매겨 근육이 부하를 담당하게 된다. 하드 캡이 아니라 비용 기반 억제이다. 이 설정의 도입 경위와 '
          '표준 설정과의 비교는 방법론적 결과이므로 3.6절에서 다룬다.',
     indent_first=0.5, align=A.JUSTIFY)
add_table(doc, 'Table 4. 동작별 실제 흡수된 척추 reserve (슈트 OFF).',
          ['동작', '척추 reserve 최대 (N·m)'],
          [[pn.NAME[k], pn.fmt(pn.reserve(k))] for k in pn.ORDER],
          widths=[6.0, 6.0], align_right_cols=(1,))
para(doc, '전 동작에서 2 N·m 이하로 균일하다. 고부하 동작에서도 tight 설정이 유효하게 작동하여 근육이 '
          '척추 부하를 담당하였다. 총 658 프레임 중 7 프레임에서 최적화가 수렴하지 않았으며, 해당 '
          '프레임을 제외하고 재계산한 결과 주 지표의 차이는 스쿼트 −0.02 %p, 스툽 0.00 %p, 박스 '
          '+0.56 %p로 결론에 영향이 없다. 보고값은 전 프레임 포함 기준이다. 동작별 실제 흡수량은 '
          'Table 4와 같다.',
     indent_first=0.5, align=A.JUSTIFY)

heading(doc, '2.7 평가 지표', 2)
para(doc, '%p는 절대 퍼센트포인트 차이, %는 상대 변화율을 뜻한다. ES peak는 각 시점에서 76개 ES 중 최대 '
          '활성 근육의 값, ES mean은 76개 평균이다. 세 가지 후보 지표를 사전 정의하고, 아래 근거로 (b)를 '
          '주 지표로 채택하였다. 선택 근거는 결과가 아니라 지표의 성질에서 도출한 것이다. 세 지표의 전 '
          '동작 값은 Table 9에 제시한다.',
     indent_first=0.5, align=A.JUSTIFY)
for t in ['(b) 채택 — 단일 순간이 아니라 구간 평균이므로 최대 활성 근육이 시점마다 바뀌어도 대표성이 유지되고, '
          '창이 슈트 토크 프로파일 자체로 객관 정의되어 결과에 의존하지 않으며, 슈트가 실제로 작동하는 구간만 평가한다.',
          '(a) 배제 — 단일 정점에 의존하여 OFF가 100 %에 포화하면 효과가 저평가되고, 최대 활성 근육이 시점마다 바뀌면 대표성을 잃는다.',
          '(c) 배제 — 76개 중 다수가 비활성이므로 평균이 희석되어 실제 부담을 과소 표현한다.']:
    para(doc, '· ' + t, align=A.JUSTIFY, space_after=3)
para(doc, '슈트 OFF 조건은 토크가 0이므로 창을 자체 정의할 수 없다. 따라서 ON 조건의 토크 프로파일에서 '
          '창을 정하고 OFF에 동일 프레임 집합을 적용하였다. 5동작 모두 OFF/ON의 시간 격자가 완전히 '
          '일치함을 확인하였다(Table 5).', indent_first=0.5, align=A.JUSTIFY)
add_table(doc, 'Table 5. 동작별 슈트 작동창 (슈트 토크 ≥ 최대치의 90 %).',
          ['동작', '창 (s)', '창 프레임 / 전체', '비고'],
          [['맨몸 스쿼트', '1.708 – 3.292', '47 / 151', '코사인 램프 (30 Hz)'],
           ['맨몸 스툽', '2.092 – 3.408', '35 / 135', '코사인 램프 (26.7 Hz)'],
           ['박스 들기', '2.225 – 5.833', '109 / 226', '하중 구간 (30 Hz)'],
           ['맨몸 보행', '전 구간', '73 / 73', '상시 24 N·m (60 Hz)'],
           ['박스 운반', '전 구간', '73 / 73', '상시 24 N·m (60 Hz)']],
          widths=[3.2, 3.6, 3.6, 3.6])

heading(doc, '2.8 동작 검증 절차', 2)
para(doc, '정량 해석 이전에 모든 동작을 렌더 기반 육안 검증으로 통과시켰다. 자가 검증의 확증 편향을 '
          '구조적으로 차단하기 위해 생성 주체와 검증 주체를 분리하였다. 검증자는 수치와 설계 의도를 '
          '제공받지 않고 렌더 이미지만으로 판정하였으며, 항목은 손이 물체를 관통하는가, 발바닥 전체가 '
          '접지하는가, 좌우가 대칭인가, 자세가 사람의 동작으로 자연스러운가이다. 미통과 동작은 해석 '
          '대상에서 제외하고 재설계하였다. 전체 해석 파이프라인을 Figure 3에 보인다.',
     indent_first=0.5, align=A.JUSTIFY)
add_figure(doc, f'{FIG}/fig2_pipeline.png',
           'Figure 3. 해석 파이프라인. 2단계(동작 육안 검증)가 관문으로, 검증 미통과 동작은 '
           '해석 대상에서 제외된다.', width_cm=16.4)

heading(doc, '2.9 해석 조건 전수 감사 및 폐기된 서술', 2)
para(doc, '본 연구는 5개 동작을 순차적으로 진행하였고, 그 과정에서 모델과 해석 설정이 갱신되었으나 앞선 '
          '동작에 소급 적용되지 않아 조건이 갈라졌다. 논문 작성 단계에서 모델 파일, SO 설정 파일, 외력 '
          '설정 파일, 결과 파일 헤더를 전수 대조하는 감사를 수행하여 이를 발견하고, 전 동작을 동일 '
          '조건으로 재해석하여 해소하였다. 재해석으로 다음 서술이 폐기되었다.',
     indent_first=0.5, align=A.JUSTIFY)
for t in ['"절대 감소량이 8.7–8.9 %p로 일정" — 통일 조건에서 부하가 걸리는 4개 동작의 절대 감소량은 −16.0 ~ −25.1 %p로 크게 다르며, 맨몸 보행은 +4.8 %p로 부호가 반대이다(Table 6).',
          '"부하–효과 단조성이 확인되어 슈트 모델이 검증됨" — 단조 경향은 주 지표에서만 관찰되고 지표 간 견고하지 않으므로 모델 검증 근거로 사용할 수 없다.',
          '"맨몸 스쿼트 47 % 감소" — 표준 reserve 조건의 값이며, 통일 조건 값은 −37.3 %이다.',
          '"정상 보행에 무영향" — 지표에 따라 부호가 갈리며 실제로는 재분배이다.']:
    para(doc, '· ' + t, align=A.JUSTIFY, space_after=3)
para(doc, '스쿼트에서 모델 변경 영향이 개별 회귀 검증의 합산 유계(약 2 %p)를 넘어 3.7 %p로 나타난 원인을 '
          '운동학 파일에서 직접 확인하였다. 쿠플러 제약은 어깨 거상각을 골반 경사의 함수로 강제한다. '
          '스툽 운동학은 이 관계를 정확히 만족하여(위반량 0.000°) 본 연구의 주 지표에서 쿠플러 제거의 '
          '영향이 −0.6 %p에 '
          '그쳤으나, 스쿼트 운동학은 양팔을 전방 85°로 규정하는 반면 쿠플러는 우측 최대 45.4°·좌측 '
          '−45.4°를 강제하여 위반량이 우측 39.6°, 좌측 130.4°에 달했다. 즉 쿠플러가 있는 상태에서는 '
          '규정된 양팔 전방 자세가 구현되지 못하고 비대칭 자세로 계산되었다. 통일 조건이 의도한 동작을 '
          '정확히 반영한 결과이며, 이는 한계가 아니라 이전 조건의 오류가 교정된 것이다.',
     indent_first=0.5, align=A.JUSTIFY)

# ============================================================ 3. 결과
heading(doc, '3. 결과', 1)
heading(doc, '3.1 동작별 ES 활성도 및 슈트 효과', 2)
_B = {k: pn.metric(k, 'b') for k in pn.ORDER}
add_table(doc, 'Table 6. 주 지표(슈트 작동창 ES peak 평균) 결과. 슈트 OFF → ON (24 N·m).',
          ['동작', '외부 하중', 'OFF (%)', 'ON (%)', 'Δ (%p)', 'Δ (%)'],
          [[pn.NAME[k], pn.LOAD[k], _B[k]['off_s'], _B[k]['on_s'],
            _B[k]['dpp_s'], '**' + _B[k]['rel_s'] + '**'] for k in pn.ORDER],
          widths=[3.0, 2.2, 2.4, 2.4, 2.4, 2.4], align_right_cols=(2, 3, 4, 5))
para(doc, '들기·운반·스툽·스쿼트에서 ES peak가 22–37 % 감소하였다. 맨몸 보행에서는 감소가 아니라 증가가 '
          '관찰되었으며, 이는 3.4절에서 재분배 관점으로 분석한다(Table 6). 동작별 시계열을 Figure 4에 '
          '보인다.',
     indent_first=0.5, align=A.JUSTIFY)
add_figure(doc, f'{FIG}/fig4_es_timeseries.png',
           'Figure 4. 동작별 척추기립근 peak 활성도 시계열. 실선은 슈트 미착용, 파선은 착용(24 N·m). '
           '음영은 (a)–(c) 슈트 작동창, (d)(e) mid-stance 구간이다.', width_cm=16.4)

heading(doc, '3.2 부하–효과 관계', 2)
para(doc, '허리에 신전 부하가 걸리는 4개 동작을 외부 하중 순으로 정렬하면, 주 지표에서 상대 감소율이 '
          '부하가 클수록 작아지는 경향이 관찰된다(Figure 5). 동일한 20 kg 하중을 갖는 두 동작(운반 '
          f'{MB["box"]["rel_s"]} %, 운반 {MB["carry"]["rel_s"]} %)이 0.6 %p 차이로 거의 일치하여, 자세 유형이 달라도 부하가 같으면 상대 '
          '효과가 유사함을 시사한다. 다만 이 경향은 주 지표에서만 성립하며 보조 지표에서는 성립하지 '
          '않는다(3.5절). 따라서 부하–효과 단조성을 본 연구의 확립된 결과로 제시하지 않으며, 모델 검증 '
          '근거로도 사용하지 않는다.', indent_first=0.5, align=A.JUSTIFY)
add_figure(doc, f'{FIG}/fig5_load_effect_pattern.png',
           'Figure 5. 슈트 24 N·m 적용 시 동작별 척추기립근 peak 활성도 변화(주 지표 기준). '
           '음수는 감소를 뜻한다. 맨몸 보행의 양수 값은 감소가 아니라 재분배를 반영한다(3.4절).',
           width_cm=12.5)

heading(doc, '3.3 절대 활성도 수준', 2)
para(doc, '통일 조건에서 슈트 미착용 시 ES peak 최대값은 스쿼트 71.7 %, 스툽 68.9 %, 박스 들기 100.0 %(포화), '
          '운반 100 %(포화), 보행 35.1 %였다. 표면 EMG 문헌이 들기 작업에서 보고하는 40–80 %MVC 범위와 '
          '부합하며, 특히 스툽 68.9 %는 선행 실측 연구[1]가 외골격 미착용 조건에서 측정한 69.8 %MVC와 '
          '0.9 %p 차이로 근접한다(4절). 20 kg 두 동작은 미착용 정점이 상한에 도달하므로 절대 부담은 '
          '보고값 이상이다(6절).', indent_first=0.5, align=A.JUSTIFY)

heading(doc, '3.4 맨몸 보행 — 감소가 아니라 재분배', 2)
para(doc, f"주 지표에서 보행의 ES peak는 {pn.metric('gait','b')['off_s']} → "
          f"{pn.metric('gait','b')['on_s']} %로 증가"
          f"(+{pn.fmt(pn.metric('gait','b')['rel'],1)} %)하는 반면, 보조 지표 ES mean은 {pn.metric('gait','c')['off_s']} → "
          f"{pn.metric('gait','c')['on_s']} %로 감소({pn.metric('gait','c')['rel_s']} %)한다. 이 상반된 결과를 근육 단위로 "
          f"분해한 결과, 활성도가 감소한 근육이 {G['n_dec']}개, 증가한 근육이 {G['n_inc']}개, 변화가 "
          f"없는 근육이 {G['n_flat']}개였다. 감소분 합은 {f(G['sum_dec'])} %p, 증가분 합은 "
          f"+{f(G['sum_inc'])} %p로 순변화는 {pn.fmt(pn.GAIT_TOTAL_DPP)} %p이며, 집중도(peak/mean)는 "
          f"{pn.fmt(pn.gait_phase(3)['conc_off'])}에서 {pn.fmt(pn.gait_phase(3)['conc_on'])}으로 상승하였다. 즉 총량은 감소하지만 부하가 소수 근육으로 집중된다. 근육군별 변화는 "
          f"Table 7에, 보행 구간별 변화는 Table 8에 정리하였다.",
     indent_first=0.5, align=A.JUSTIFY)
_GG = [('장늑근 (Iliocostalis)', 'Iliocostalis (IL)'),
       ('최장근 요추부 (LTpL)', 'Longissimus pars lumborum (LTpL)'),
       ('최장근 흉추부 (LTpT)', 'Longissimus pars thoracis (LTpT)')]
add_table(doc, 'Table 7. 맨몸 보행의 근육군별 재분배 (전주기 시간평균 활성도 합).',
          ['근육군', '개수', 'OFF', 'ON', '변화'],
          [[lbl, str(pn.gait_group(key)['n']), pn.gait_group(key)['off_s'],
            pn.gait_group(key)['on_s'],
            ('**' + pn.gait_group(key)['rel_s'] + ' %**') if abs(pn.gait_group(key)['rel']) > 10
            else (pn.gait_group(key)['rel_s'] + ' %')] for lbl, key in _GG],
          widths=[5.4, 1.8, 2.4, 2.4, 3.0], align_right_cols=(1, 2, 3, 4))
_PH = [pn.gait_phase(i) for i in range(len(G['phases']))]
add_table(doc, 'Table 8. 맨몸 보행의 구간별 재분배.',
          ['구간', 'ES peak OFF → ON', 'ES mean OFF → ON', '집중도 OFF → ON'],
          [[r['phase'],
            f"{pn.fmt(r['peak_off'])} → {pn.fmt(r['peak_on'])} ({pn.fmt(r['peak_rel'],1,True)} %)",
            f"{pn.fmt(r['mean_off'])} → {pn.fmt(r['mean_on'])} ({pn.fmt(r['mean_rel'],1,True)} %)",
            f"{pn.fmt(r['conc_off'])} → {pn.fmt(r['conc_on'])}"] for r in _PH],
          widths=[2.6, 4.6, 4.6, 3.2])
para(doc, '세 구간 모두에서 최대 활성 근육은 증가하고 평균은 감소하며 집중도가 상승한다. 재분배는 특정 '
          '보행 국면에 국한되지 않고 주기 전반에 걸쳐 나타난다.', indent_first=0.5, align=A.JUSTIFY)
para(doc, '기전을 세 동작 대조로 검토하였다. 최장근 요추부 증가는 보행에서만 나타나며, 저부하만으로 '
          '설명되지 않고(스툽도 0 kg이나 전 근육군 감소), 면외 운동만으로도 설명되지 않는다(운반도 동일 '
          '운동학이나 전 근육군 감소). 운동학을 확인한 결과 보행은 흉요추 축회전 진폭 11.63°, 측굴 '
          '10.87°를 갖는 반면 스툽은 두 값이 정확히 0.00°였다. 보행은 시상면 신전 요구가 작아 상시 '
          '24 N·m가 이를 과충족시키고, 그 결과 시상면 신전을 주로 담당하던 장늑근이 거의 완전히 '
          '비활성화된다. 그러나 축회전·측굴의 면외 요구는 그대로 남으며, 장늑근이 부수적으로 담당하던 '
          '몫이 최장근 요추부로 이전되어 증가한다. 이 해석은 활성도 패턴과 운동학에서 도출한 것이며, '
          '근육별 모멘트 암 분해로 직접 확인한 것은 아니다.', indent_first=0.5, align=A.JUSTIFY)
para(doc, '총량이 감소한 것은 전체 부하 경감으로 볼 여지가 있으나, 동시에 특정 근육군에 부하가 집중되는 '
          '것은 국소 피로·손상 위험 증가로 해석될 여지도 있다. 본 연구의 데이터만으로는 어느 쪽인지 '
          '단정할 수 없으며 실측 EMG 검증이 필요하다. 다만 변화의 절대 크기는 들기 동작(−16 ~ −25 %p) '
          '대비 작다(+4.79 %p). 근육별 재분배 양상을 Figure 6에 보인다.',
     indent_first=0.5, align=A.JUSTIFY)
add_figure(doc, f'{FIG}/fig7_gait_redistribution.png',
           'Figure 6. 맨몸 보행 조건의 근육별 재분배. (a) 최대 활성 근육은 증가, (b) 76근육 평균은 감소, '
           '(c) 근육별 시간평균 활성도 변화 분포(감소 37개·증가 10개, 변화량 오름차순 정렬).',
           width_cm=16.4)

heading(doc, '3.5 주 지표 선택이 결론에 미치는 영향', 2)
def _cell(k, m):
    d = pn.metric(k, m)
    return f"{d['off_s']} → {d['on_s']}\n({d['rel_s']} %)"
add_table(doc, 'Table 9. 지표 3종 전 동작 대조. 괄호 안은 상대 변화율.',
          ['동작', '하중', '(a) 전주기 peak', '(b) 슈트 작동창 peak 평균 (주 지표)', '(c) ES mean 창평균'],
          [[pn.NAME[k], pn.LOAD[k], _cell(k, 'a'), '**' + _cell(k, 'b') + '**', _cell(k, 'c')]
           for k in pn.ORDER]
          + [['부하 순 단조 경향\n(부하 있는 4개 동작 기준)', '—', '성립하지 않음', '**성립**', '성립하지 않음']],
          widths=[2.6, 1.5, 3.9, 4.5, 3.9], size=8.6)
para(doc, f'지표 간 편차의 원인은 각 지표의 정의에서 설명된다. 박스 운반의 (a) {MA["carry"]["rel_s"]} %는 슈트 미착용에서 '
          '최대 활성 근육이 100 %에 도달하여 더 이상 커질 수 없기 때문이며, 정점값이 실제 요구를 '
          f'반영하지 못해 효과가 저평가된다. 창 평균 지표에서는 {MB["carry"]["rel_s"]} %로 회복된다. '
          f'박스 들기도 미착용 정점이 {MA["box"]["off_s"]} %로 같은 성질을 갖는다. 박스 운반의 (a) '
          f'{MA["carry"]["rel_s"]} %와 (c) {MC["carry"]["rel_s"]} % 차이는 22.6 %p에 달하는데, 정점은 포화로 눌리고 평균은 '
          '까지 변동하기 때문이며, 최대 활성 근육이 시점마다 바뀌는 데서 비롯된다(활성도 90 % 이상 '
          '프레임은 0개로 포화가 아니다). (c)는 76개 중 다수가 비활성이어서 평균이 희석되며, 보행에서 '
          '이 성질이 극단적으로 나타나 미착용 baseline이 1.92 %에 불과하고 상대 변화율이 불안정해진다.',
     indent_first=0.5, align=A.JUSTIFY)
para(doc, '동일한 시뮬레이션 결과에서 지표 정의만 바꾸어도 개별 동작의 감소율이 최대 10 %p 이상 '
          '달라지고, 보행에서는 부호까지 반전된다. 웨어러블 보조 효과를 시뮬레이션으로 평가하는 '
          '연구에서 지표 정의는 부수적 선택이 아니라 결론을 좌우하는 핵심 설계 요소이다. 후속 연구에는 '
          '지표를 결과 확인 전에 사전 정의할 것, 포화와 근육 전환 가능성을 사전 점검할 것, 단일 지표가 '
          '아니라 복수 지표를 병기할 것을 권고한다.', indent_first=0.5, align=A.JUSTIFY)

heading(doc, '3.6 Reserve 설정 민감도', 2)
_S, _T = pn.RES_SENS['std'], pn.RES_SENS['tight']
add_table(doc, 'Table 10. 맨몸 보행 조건의 reserve 설정 민감도.',
          ['항목', '표준 reserve (척추 opt 100 N·m)', 'tight reserve (척추 opt 5 N·m)'],
          [['실제 흡수된 척추 reserve 최대', f"{pn.fmt(_S['reserve'])} N·m", f"**{pn.fmt(_T['reserve'])} N·m**"],
           ['보행 ES peak, 슈트 OFF', f"{pn.fmt(_S['off'])} %", f"**{pn.fmt(_T['off'])} %**"],
           ['보행 ES peak, 슈트 ON', f"{pn.fmt(_S['on'])} %", f"{pn.fmt(_T['on'])} %"],
           ['슈트 효과 (전주기 peak)', f"**{_S['dpp_s']} %p (\"보조\")**", f"**{_T['dpp_s']} %p**"]],
          widths=[5.4, 5.3, 5.3], align_right_cols=(1, 2))
para(doc, '표준 설정에서는 reserve가 척추 신전 부하를 근육 대신 흡수하여 ES 활성도를 약 3.2배 '
          '과소평가하였다. 더 중요하게는 reserve가 흡수하는 부하량이 슈트 착용 시 함께 감소하면서 '
          '실재하지 않는 보조 효과 −5.6 %p를 만들어냈다. 또한 표준 설정에서는 스쿼트·스툽·박스의 절대 '
          '활성도가 각각 23.1 / 31.9 / 37.5 %로 EMG 문헌 범위에 크게 미달했으나, tight 설정에서 '
          '71.7 / 68.9 / 100.0 %로 문헌 범위 이상으로 올라왔다. 저부하 동작을 해석할 때는 reserve 크기와 '
          '실제 활성도를 반드시 점검해야 하며, 점검 없이는 결론의 부호까지 뒤바뀔 수 있다. '
          'reserve 설정에 따른 차이를 Table 10과 Figure 7에 보인다.',
     indent_first=0.5, align=A.JUSTIFY)
add_figure(doc, f'{FIG}/fig6_reserve_sensitivity.png',
           'Figure 7. reserve 설정에 따른 척추기립근 활성도 추정과 슈트 효과의 차이. '
           '(a) reserve가 흡수한 척추 부하, (b) 근육 활성도와 슈트 효과 추정.', width_cm=16.4)

# ============================================================ 4. 문헌
heading(doc, '4. 선행 연구 대조', 1)
para(doc, '본 절의 수치는 모두 원저 초록 원문에서 조건·지표·수치를 직접 확인한 것만 기재하였다.',
     indent_first=0.5, align=A.JUSTIFY)
para(doc, 'Hasenmaier 등[1]은 건강한 젊은 성인 17명(근전도 분석 16명)을 대상으로 Apogee 능동 외골격 '
          '착용 하에 대칭 stoop 및 squat 들기를 수행하며 척추기립근 표면 근전도를 측정하였다. stoop '
          '기법에서 미착용 69.8 %MVC, 0/0 % 보조 59.2 %, 50/20 % 보조 50.7 %, 100/60 % 보조 42.4 %를 '
          '보고하였다. 원문은 두 기법 전체에 대해 "약 10–27 %MVC의 감소"로 요약하는데, 이는 %MVC 절대 '
          '포인트 감소이며 상대 감소율이 아니다. 미착용에서 최대 보조까지의 상대 감소율은 −39.3 %이다. '
          'squat 기법에 대해서는 보조 수준 간 유의차가 없다고 보고하여 상대 감소율을 인용할 수 없다.',
     indent_first=0.5, align=A.JUSTIFY)
para(doc, f'본 연구의 통일 조건 스툽 미착용 ES peak 최대값은 {MA["stoop"]["off_s"]} %로 Hasenmaier 등의 69.8 %MVC와 '
          '0.9 %p 차이이다. 표준 reserve 조건의 값(31.9 %)은 문헌 범위에 크게 미달했으므로, 이 정합은 '
          f'tight reserve 설정의 타당성을 뒷받침하는 독립적 근거이다. 슈트 효과는 본 연구 {MB["stoop"]["rel_s"]} %(주 '
          f'지표) 또는 {MA["stoop"]["rel_s"]} %(전주기 peak)로 Hasenmaier 등의 −39.3 %보다 작아, 본 시뮬레이션이 슈트 '
          '효과를 과대평가하지 않았음을 시사한다. 다만 본 연구의 ES peak는 76개 근육 중 최대값이고 '
          'Hasenmaier 등의 값은 표면 전극이 포착하는 표층 근육의 %MVC이므로, 근육 선정 범위와 정규화 '
          '기준이 다르다. 위 정합은 엄밀한 등치가 아니라 크기 수준의 대조로 해석해야 하며, 보조 토크 '
          '사양도 다르다.', indent_first=0.5, align=A.JUSTIFY)
para(doc, 'Hu 등[2]은 8명을 대상으로 15 kg 자유 기법 들기에서 능동 이중관절 등 지지 외골격의 효과를 '
          '보고하였다. 등 근육 능동 모멘트 −14.9 ~ −28.6 %, L5/S1 압축력 −5.5 ~ −9.3 %이다. 두 값 모두 '
          '관절 부하 지표이며 근육 활성도가 아니므로 본 연구의 ES 활성도 감소와 직접 비교할 수 없고, '
          '관절 부하 지표에서 관찰된 감소 경향과 방향성이 일치한다는 수준으로만 대조한다. 대조 결과를 '
          'Table 11에 정리하였다.',
     indent_first=0.5, align=A.JUSTIFY)
add_table(doc, 'Table 11. 선행 연구 대조. 보고 지표가 연구마다 다름에 유의.',
          ['출처', '지표', '조건', '보고값', '본 연구 대응값'],
          [['Hasenmaier 2026', 'ES 활성도 (%MVC)', 'stoop, 미착용', '69.8', f'스툽 OFF {MA["stoop"]["off_s"]} % (정합)'],
           ['Hasenmaier 2026', 'ES 활성도', 'stoop, 100/60 % 보조', '42.4 (상대 −39.3 %)', f'스툽 {MB["stoop"]["rel_s"]} % (더 보수적)'],
           ['Hasenmaier 2026', 'ES 활성도', 'squat', '수준 간 유의차 없음', '스쿼트 −37.3 % — 대조 불가'],
           ['Hu 2026', '등 근육 능동 모멘트', '15 kg 들기', '−14.9 ~ −28.6 %', '지표 상이 — 방향성만 대조'],
           ['Hu 2026', 'L5/S1 압축력', '15 kg 들기', '−5.5 ~ −9.3 %', '본 연구 미산출']],
          widths=[3.0, 3.2, 3.0, 3.0, 3.8])

# ============================================================ 5. 고찰
heading(doc, '5. 고찰', 1)
heading(doc, '5.1 고정 보조 토크와 부하 의존성', 2)
para(doc, '슈트 보조 토크는 24 N·m로 고정되어 있으므로 허리 요구 토크가 커질수록 상대 비중이 감소한다. '
          '주 지표에서 관찰된 부하 의존 경향은 이 구조와 정합한다. 다만 이 경향은 지표 간 견고하지 '
          '않으므로 설계 지침으로 활용할 때는 지표 정의를 함께 명시해야 한다. 실무적으로는, 고부하 '
          '작업에서 더 큰 비율의 효과를 목표로 한다면 보조 토크 자체를 증대해야 하며, 중·저부하 작업이 '
          '주 사용 시나리오라면 현 24 N·m 수준으로도 상당한 상대 효과를 얻을 수 있다.',
     indent_first=0.5, align=A.JUSTIFY)
heading(doc, '5.2 절대값과 상대값의 견고성 차이', 2)
para(doc, '조건 통일 과정에서 절대 활성도는 조건에 민감하고 상대 지표는 상대적으로 견고함이 드러났다. '
          '기저 모델만 바꾸어도 슈트 미착용 baseline이 최대 +7.72 %p 이동한 반면(스쿼트), 주 지표의 '
          '상대 변화는 0.6–3.7 %p에 그쳤다. reserve 설정 변경 시에는 절대값이 스쿼트·스툽·박스에서 '
          '2.2–3.1배, 보행에서 3.2배 변했으나 상대 '
          '지표는 그보다 훨씬 작게 움직였다. 따라서 절대 활성도를 인용할 때는 모델과 reserve 조건을 '
          '반드시 병기해야 하며, 조건이 다른 연구 간 절대값 비교는 신중해야 한다. 상대 지표는 조건 '
          '변화에 더 견고하나 지표 정의 자체에는 민감하다.', indent_first=0.5, align=A.JUSTIFY)
heading(doc, '5.3 선택적 보조와 상시 구동의 트레이드오프', 2)
para(doc, f'맨몸 보행과 박스 운반은 운동학이 유사하고 하중만 다르다. 운반에서는 {MB["carry"]["rel_s"]} %의 감소가 '
          '나타나지만 보행에서는 총량 감소와 함께 최대 활성 근육의 증가라는 재분배가 나타난다. SMA '
          '액추에이터는 가열–냉각을 반복하는 On/Off 구동보다 50 ℃를 유지하는 상시 구동이 잠열 재투입을 '
          '피할 수 있어 에너지 효율이 약 13배 높다. 그러나 보행 중 상시 신전 토크가 부하를 줄이는 대신 '
          '최심부 요추 신전근으로 옮긴다면, 에너지 효율의 이점과 보행 적합성 사이에 트레이드오프가 '
          '존재할 수 있다. 이로부터 보행 구간에서 토크를 저감하거나 차단하는 제어가 유리한지 검토할 '
          '필요가 제기된다. 다만 이는 본 연구가 검증한 결론이 아니라 본 연구 결과로부터 제기하는 '
          '가설이며, 재분배가 실제로 유해한지 여부는 실측 근전도로 확인해야 한다.',
     indent_first=0.5, align=A.JUSTIFY)
heading(doc, '5.4 방법론적 시사점', 2)
for t in ['저부하 조건에서 reserve 설정은 결론을 좌우한다. 표준 설정은 ES를 3.2배 과소평가하고 실재하지 않는 보조 효과를 만들어냈다.',
          '지표 정의가 결론을 좌우한다. 동일 결과에서 지표만 바꾸어도 감소율이 10 %p 이상, 보행에서는 부호까지 달라진다.',
          '다동작 비교 연구는 조건 통일을 사전에 강제해야 한다. 매니페스트 기반 조건 관리와 자동 점검을 권고한다.',
          '동작의 물리적 타당성 검증이 정량 신뢰도의 전제이다. 생성 주체와 검증 주체를 분리하는 절차가 자가 검증의 확증 편향을 차단한다.']:
    para(doc, '· ' + t, align=A.JUSTIFY, space_after=3)

# ============================================================ 6. 한계
heading(doc, '6. 한계', 1)
for t in ['성인 남성 1개 체형만 해석하였다. 주 적용 대상으로 상정한 고령 간병 인력에 대한 정량값은 별도 확장이 필요하다.',
          '박스 들기는 지면반력을 명시적으로 부여하지 않고 골반 reserve가 체중 지지력을 흡수한다. 슈트 효과 차이는 견고하나 절대 활성도의 동작 간 직접 비교에는 이 차이를 고려해야 한다.',
          '보행·운반의 지면반력은 다른 피험자의 실측값을 체중 스케일로 보정하여 사용하였다. 슈트 ON/OFF가 동일 지면반력을 공유하므로 차이는 견고하다.',
          'Static Optimization 기반이므로 근육 활성 동역학과 길이–속도 의존성이 반영되지 않으며, 각 시점을 독립적으로 해결하여 시간적 연속성이 보장되지 않는다.',
          '운반 조건에서 슈트 미착용 시 최대 활성 근육이 100 %에 도달하여 실제 부담은 보고값 이상이다. 주 지표는 이를 부분적으로 완화하나 완전히 해소하지 못한다.',
          '스쿼트 조건은 대응 선행 연구가 유의차를 보고하지 못하여 외부 검증 없이 제시된다. 실측 근전도 검증의 최우선 대상이다.',
          '보행 재분배의 유해성 여부는 본 연구 데이터로 판단할 수 없다. 재분배 기전의 해석도 활성도 패턴과 운동학에서 도출한 것이며 모멘트 암 분해로 직접 확인하지 않았다.',
          '총 658 프레임 중 7개에서 최적화가 수렴하지 않았다. 주 지표에 미치는 영향은 0.56 %p 이하이다.',
          '슈트를 흉추–골반 간 순수 토크 커플로 이상화하였다. 실제 착용 시의 의복 슬립, 연부조직 변형, 착용 위치 편차에 의한 유효 모멘트 암 감소는 반영되지 않았으며, 실제 효과를 과대평가하는 방향으로 작용할 수 있다.',
          '좌측 어깨 거상 자유도의 축 정의에 잔여 문제가 있다. 본 5동작은 영향권 밖이나 어깨 보조를 포함하는 후속 해석 전에 정량 진단이 필요하다.']:
    para(doc, '· ' + t, align=A.JUSTIFY, space_after=3)

# ============================================================ 7. 향후
heading(doc, '7. 향후 과제', 1)
for t in ['실측 근전도 검증 — 특히 스쿼트 조건과 보행 재분배의 유해성 판단',
          '보행 구간 토크 제어 검토 — 5.3절에서 제기한 트레이드오프 가설의 검증',
          '모멘트 암 분해 — 보행 재분배 기전의 직접 확인',
          '다관절 보조 효과 평가 — 허리에 더해 양측 어깨·팔꿈치',
          '인체 조건 확장 — 성별·연령별 근력 및 인체 비율 조정 모델',
          '동적 해석 — OpenSim Moco 기반 활성 동역학 반영',
          'L5/S1 압축력 산출 — 선행 연구와 동일 지표로 대조 가능하게 함']:
    para(doc, '· ' + t, align=A.JUSTIFY, space_after=3)

# ============================================================ 8. 결론
heading(doc, '8. 결론', 1)
for i, t in enumerate([
    'SMA 직물 근육 기반 허리 보조 슈트(24 N·m)는 5동작 완전 통일 조건에서 척추기립근 peak 활성도(슈트 '
    f'작동창 평균)를 맨몸 스쿼트 {MB["squat"]["rel_s"]} %, 맨몸 스툽 {MB["stoop"]["rel_s"]} %, '
    f'박스 들기 {MB["box"]["rel_s"]} %, 박스 운반 {MB["carry"]["rel_s"]} % '
    '변화시켰다.',
    '주 지표에서는 부하가 클수록 상대 감소율이 작아지는 경향이 관찰되나 이 경향은 지표 간 견고하지 '
    '않다. 지표 정의만 바꾸어도 개별 동작 감소율이 10 %p 이상 달라지므로, 지표 정의는 결론을 좌우하는 '
    '핵심 설계 요소이다.',
    '정상 보행에서는 부하 감소가 아니라 재분배가 관찰되었다. 총량은 −17.4 %p 감소하나 최대 활성 근육은 '
    '+21.4 % 증가하며, 장늑근이 −90.9 % 비활성화되는 대신 최장근 요추부가 +28.5 % 증가한다. 유해성 '
    '여부는 실측 검증이 필요하다.',
    '방법론적으로, 저부하 동작 해석에서 reserve 설정이 근육 활성도를 3.2배 과소평가하고 슈트 효과의 '
    '존재 여부까지 왜곡할 수 있음을 제시하였다.',
    '통일 조건의 절대 활성도(스툽 미착용 68.9 %)가 실측 근전도 문헌(69.8 %MVC)과 0.9 %p 차이로 '
    '정합하며, 슈트 효과는 실측치보다 보수적이다.']):
    para(doc, f'({i+1}) ' + t, align=A.JUSTIFY, space_after=4)

# ============================================================ 참고문헌
heading(doc, '참고문헌', 1)
for t in [
    '[1] Hasenmaier J, Siebert T, Mayer D, Stutzig N. Effects of an active exoskeleton on the muscle '
    'activity of the erector spinae and biceps femoris muscles during lifting with symmetric stoop and '
    'squat technique. Frontiers in Bioengineering and Biotechnology, 2026. doi:10.3389/fbioe.2026.1631785',
    '[2] Hu F, Brouwer NP, Tabasi A, et al. Influence of varied assistance levels provided by a '
    'dual-joint active back-support exoskeleton on spinal musculoskeletal loading and kinematics during '
    'lifting. Ergonomics, 2026;69(3):453–465. doi:10.1080/00140139.2025.2466030 (PMID 39967340)',
    '[3] Beaucage-Gauvreau E, Robertson WSP, Brandon SCE, et al. Validation of an OpenSim full-body '
    'model with detailed lumbar spine for estimating lower lumbar spine loads during symmetric and '
    'asymmetric lifting tasks. Computer Methods in Biomechanics and Biomedical Engineering, '
    '2019;22(5):451–464.']:
    para(doc, t, size=9.5, line=1.45, align=A.JUSTIFY, space_after=5)

doc.save(OUT)
print('SAVED', OUT, round(os.path.getsize(OUT) / 1e6, 2), 'MB')
